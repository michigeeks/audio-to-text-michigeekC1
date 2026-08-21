import os
import tempfile
from datetime import datetime
from io import BytesIO

import streamlit as st
from docx import Document
from faster_whisper import WhisperModel


# --------------------------------------------------
# Configuración de página
# --------------------------------------------------

st.set_page_config(
    page_title="Audio → Texto - Michigeeks v2",
    page_icon="🎙️",
    layout="wide",
)


# --------------------------------------------------
# Estado inicial
# --------------------------------------------------

if "texto" not in st.session_state:
    st.session_state.texto = ""
if "duracion" not in st.session_state:
    st.session_state.duracion = None
if "nombre_audio" not in st.session_state:
    st.session_state.nombre_audio = None


# --------------------------------------------------
# Configuración hardcodeada
# --------------------------------------------------

TAMAÑO_MODELO = "medium"
IDIOMA = "es"

MIME_POR_EXTENSION = {
    ".mp3": "audio/mp3",
    ".wav": "audio/wav",
    ".m4a": "audio/mp4",
    ".ogg": "audio/ogg",
    ".flac": "audio/flac",
    ".aac": "audio/aac",
}
# --------------------------------------------------
# Cargar modelo
# --------------------------------------------------

@st.cache_resource(show_spinner=False)
def cargar_modelo():
    return WhisperModel(TAMAÑO_MODELO, device="cpu", compute_type="int8")


# --------------------------------------------------
# Utilidades
# --------------------------------------------------

def formatear_tiempo(segundos):
    m, s = divmod(int(segundos), 60)
    h, m = divmod(m, 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def generar_docx(texto, nombre_audio):
    doc = Document()

    doc.add_heading("Transcripción de audio", level=1)

    meta = doc.add_paragraph()
    meta.add_run("Archivo: ").bold = True
    meta.add_run(nombre_audio or "—")
    meta.add_run("\nFecha: ").bold = True
    meta.add_run(f"{datetime.now():%d/%m/%Y %H:%M}")

    doc.add_paragraph()  # espacio
    doc.add_paragraph(texto)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --------------------------------------------------
# Header
# --------------------------------------------------

st.title("🎙️ Audio a Texto")
st.markdown("Convierte archivos de audio a texto automáticamente.")

st.divider()

col1, col2 = st.columns(2, gap="large")


# --------------------------------------------------
# Columna 1 — Audio
# --------------------------------------------------

with col1:
    st.subheader("🎧 Audio de entrada")

    audio = st.file_uploader(
        "Sube o arrastra un archivo de audio",
        type=["mp3", "wav", "m4a", "ogg", "flac", "aac"],
        help="Tamaño máximo recomendado: 200 MB",
    )

    if audio is not None:
        ext = os.path.splitext(audio.name)[1].lower()
        mime = MIME_POR_EXTENSION.get(ext, "audio/wav")
        st.audio(audio.getvalue(), format=mime)

        tamaño_mb = audio.size / (1024 * 1024)
        st.caption(f"📁 {audio.name} · {tamaño_mb:.1f} MB")

        if tamaño_mb > 200:
            st.warning("⚠️ El archivo es grande, la transcripción puede demorar bastante.")

        transcribir_btn = st.button("🚀 Transcribir audio", type="primary", use_container_width=True)
    else:
        transcribir_btn = False
        st.info("👆 Sube un archivo de audio para comenzar.")


# --------------------------------------------------
# Columna 2 — Resultado
# --------------------------------------------------

with col2:
    st.subheader("📝 Transcripción")

    if transcribir_btn:
        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(audio.name)[1]) as temp:
                temp.write(audio.getbuffer())
                temp_path = temp.name

            with st.status("Procesando audio...", expanded=True) as status:
                st.write("🔊 Cargando audio...")
                model = cargar_modelo()

                st.write("🧠 Transcribiendo...")
                segments_gen, info = model.transcribe(temp_path, language=IDIOMA, vad_filter=True)

                progress_bar = st.progress(0, text="Procesando segmentos...")
                segmentos = []
                duracion_total = info.duration or 1

                for seg in segments_gen:
                    segmentos.append(seg)
                    avance = min(seg.end / duracion_total, 1.0)
                    progress_bar.progress(avance, text=f"Procesando... {formatear_tiempo(seg.end)} / {formatear_tiempo(duracion_total)}")

                progress_bar.empty()

                texto = " ".join(s.text.strip() for s in segmentos)

                st.session_state.texto = texto
                st.session_state.duracion = duracion_total
                st.session_state.nombre_audio = audio.name

                status.update(label="✅ Transcripción completada", state="complete", expanded=False)

            st.toast("Transcripción lista 🎉")

        except Exception as e:
            st.error(f"❌ Ocurrió un error al transcribir: {e}")
            st.session_state.texto = ""

        finally:
            if temp_path and os.path.exists(temp_path):
                os.remove(temp_path)

    # --- Mostrar resultados (persisten entre reruns gracias a session_state) ---
    if st.session_state.texto:

        st.metric("⏱️ Duración del audio", formatear_tiempo(st.session_state.duracion))

        st.text_area("Resultado", value=st.session_state.texto, height=300)

        nombre_base = os.path.splitext(st.session_state.nombre_audio or "transcripcion")[0]
        docx_buffer = generar_docx(st.session_state.texto, st.session_state.nombre_audio)
        st.download_button(
            "⬇️ Descargar como Word 📄 (.docx)",
            data=docx_buffer,
            file_name=f"{nombre_base}_{datetime.now():%Y%m%d_%H%M}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    elif not transcribir_btn:
        st.info("Sube un audio y pulsá «Transcribir» para ver el resultado acá.")
